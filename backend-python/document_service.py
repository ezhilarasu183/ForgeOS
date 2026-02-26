import json
import os
import sqlite3
from datetime import date
from groq import Groq
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import threading

class DocumentService:
    def __init__(self, db_path, config_file):
        self.db_path = db_path
        self.config_file = config_file
        self.api_key = self._load_api_key()
        self.client = Groq(api_key=self.api_key) if self.api_key else None
        
        # Ensure uploads directory exists
        self.upload_dir = os.path.join(os.path.dirname(db_path), "uploads")
        if not os.path.exists(self.upload_dir):
            os.makedirs(self.upload_dir)

    def _load_api_key(self):
        try:
            with open(self.config_file, 'r') as f:
                config = json.load(f)
                return config.get("groq_key")
        except:
            return None

    def _get_db_connection(self):
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        return conn

    def understand_agent(self, user_input):
        if not self.client: return None
        prompt = """extract structured information from the request
        Important rules:
            - Do NOT assume missing information.
            - If something is not explicitly stated, return null.
            - Do NOT infer architecture or deployment.
            - Only extract what is clearly mentioned.
            return only valid JSON with these keys {
          "ProjectName",
          "Purpose",
          "Scope",
          "Stakeholders",
          "FunctionalRequirements",
          "NonFunctionalRequirements",
          "Constraints",
          "AcceptanceCriteria",
          "UseCases"
        }
        user prompt : """ + user_input
        
        response = self.client.chat.completions.create(
                model="llama-3.3-70b-versatile", 
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000,
                temperature=0.0,
            )
        raw_text = response.choices[0].message.content.strip()
        start = raw_text.find("{")
        end = raw_text.rfind("}") + 1
        if start == -1 or end == 0: return None
        return json.loads(raw_text[start:end])

    def generate_enhanced_content(self, final_data):
        if not self.client: return None
        format_template = {
            "ProjectName": "", "Purpose": "", "Scope": "", "Stakeholders": [],
            "FunctionalRequirements": [], "NonFunctionalRequirements": [],
            "Constraints": [], "AcceptanceCriteria": [], "UseCases": []
        }
        
        prompt = f"""
You are a professional System Requirements Engineer.
Return ONLY valid JSON.
STRUCTURE (follow exactly):
{json.dumps(format_template, indent=2)}
INPUT DATA:
{json.dumps(final_data, indent=2)}
CONTENT RULES:
1. Purpose & Scope: Professional paragraphs (4-5 lines).
2. All list fields: JSON arrays, 6-9 professional sentences each. No subheadings.
"""
        response = self.client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2000,
            temperature=0.0,
        )
        raw_text = response.choices[0].message.content.strip()
        start = raw_text.find("{")
        end = raw_text.rfind("}") + 1
        return json.loads(raw_text[start:end])

    def _robust_replace(self, target, placeholder, value):
        """Unified replacement for paragraphs and table cells that handles split runs."""
        # This handles paragraphs and cells (which have .paragraphs)
        for p in target.paragraphs:
            if placeholder in p.text:
                # Basic replacement
                # Note: assignment to p.text preserves paragraph style but resets run formatting
                # However, for simple placeholders, this is often the most reliable way 
                # to ensure we don't miss split runs. We'll try to re-apply basic formatting.
                inline = p.runs
                # Try to capture the formatting of the first run if it exists
                font_name = "Inter"
                font_size = Pt(11)
                if inline:
                    font_name = inline[0].font.name
                    font_size = inline[0].font.size

                new_text = p.text.replace(placeholder, str(value))
                p.text = "" # Clears runs
                run = p.add_run(new_text)
                if font_name: run.font.name = font_name
                if font_size: run.font.size = font_size

    def _apply_paragraphs(self, placeholder, value, doc):
        # 1. Body paragraphs
        self._robust_replace(doc, placeholder, value)
        
        # 2. Table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._robust_replace(cell, placeholder, value)

    def _insert_bullets(self, placeholder, items, doc):
        if not items: return
        
        def process_container(container):
            for para in container.paragraphs:
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, "")
                    current_p = para
                    for item in items:
                        try:
                            # Try to use 'List Bullet' if it exists in template
                            new_p = doc.add_paragraph(str(item), style='List Bullet')
                        except:
                            new_p = doc.add_paragraph("• " + str(item))
                        
                        # Set font to match (optional but keeps it clean)
                        if new_p.runs:
                            new_p.runs[0].font.name = "Inter"
                            new_p.runs[0].font.size = Pt(11)
                            
                        current_p._element.addnext(new_p._element)
                        current_p = new_p
                    return True
            return False

        # Body
        if process_container(doc): return
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if process_container(cell): return

    def ssd_json(self, data, bom="N/A"):
        if not self.client: return None
        format_spec = {
            "ProjectName": "", "Purpose": "", "Scope": "", "SystemArchitecture": [],
            "HardwareComponentSpecification": [], "SoftwareStack": [], "Interfaces": [],
            "ElectricalMechanicalIntegration": [], "SafetyFailover": [], "Constraints": []
        }
        prompt = f"""You are a senior System Architect. Return ONLY valid JSON.
{json.dumps(format_spec, indent=2)}
INPUT SRD: {json.dumps(data, indent=2)}
INPUT BOM: {bom}
"""
        response = self.client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2000,
            temperature=0.0,
        )
        ssd_data_raw = response.choices[0].message.content.strip()
        start = ssd_data_raw.find("{")
        end = ssd_data_raw.rfind("}") + 1
        return json.loads(ssd_data_raw[start:end])

    def generate_project_docs(self, project_id):
        conn = self._get_db_connection()
        project = conn.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
        if not project:
            conn.close()
            return
            
        description = project['description'] or project['name']
        data = self.understand_agent(description)
        if not data:
            data = {"ProjectName": project['name'], "Purpose": description, "Scope": "", "Stakeholders": [], "FunctionalRequirements": [], "NonFunctionalRequirements": [], "Constraints": [], "AcceptanceCriteria": [], "UseCases": []}
            
        enhanced_data = self.generate_enhanced_content(data) or data
        author = project['owner'] or "AI Assistant"
        today = date.today().strftime("%Y-%m-%d")
        
        # Paths
        srd_filename = f"SRD_{project_id}.docx"
        ssd_filename = f"SSD_{project_id}.docx"
        srd_path = os.path.join(self.upload_dir, srd_filename)
        ssd_path = os.path.join(self.upload_dir, ssd_filename)
        
        # templates - search upwards from backend directory
        current_search_dir = os.path.dirname(self.db_path)
        srd_template = None
        ssd_template = None
        
        # Search up to 5 levels for templates
        for _ in range(5):
            potential_srd = os.path.join(current_search_dir, "srd_template.docx")
            potential_ssd = os.path.join(current_search_dir, "ssd_template.docx")
            
            if not srd_template and os.path.exists(potential_srd):
                srd_template = potential_srd
            if not ssd_template and os.path.exists(potential_ssd):
                ssd_template = potential_ssd
                
            if srd_template and ssd_template: break
            
            parent = os.path.dirname(current_search_dir)
            if parent == current_search_dir: break # reached root
            current_search_dir = parent

        # Fallback to current dir if not found (though srd_template and ssd_template will be None)
        if not srd_template: srd_template = "srd_template.docx"
        if not ssd_template: ssd_template = "ssd_template.docx"
        
        srd_generated = False
        ssd_generated = False
        
        # Generate SRD
        if os.path.exists(srd_template):
            try:
                doc1 = Document(srd_template)
                placeholders = {
                    "{{PROJECT NAME}}": enhanced_data['ProjectName'],
                    "{{PURPOSE}}": enhanced_data['Purpose'],
                    "{{SCOPE}}": enhanced_data['Scope'],
                    "{{VERSION}}": "1.0",
                    "{{DATE}}": today,
                    "{{AUTHOR}}": author
                }
                for k, v in placeholders.items(): self._apply_paragraphs(k, v, doc1)
                
                bullets = {
                    "{{STAKEHOLDERS}}": "Stakeholders",
                    "{{HIGH-LEVEL SYSTEM REQUIREMENTS (FUNCTIONAL)}}": "FunctionalRequirements",
                    "{{NON-FUNCTIONAL REQUIREMENTS}}": "NonFunctionalRequirements",
                    "{{CONSTRAINTS}}": "Constraints",
                    "{{ACCEPTANCE CRITERIA}}": "AcceptanceCriteria",
                    "{{USECASE}}": "UseCases"
                }
                for k, v in bullets.items(): self._insert_bullets(k, enhanced_data[v], doc1)
                
                doc1.save(srd_path)
                srd_generated = True
                print(f"SRD generated: {srd_path}")
            except Exception as e:
                print(f"SRD generation error: {e}")
        else:
            print(f"SRD template missing at {srd_template}")
        
        # Generate SSD
        if os.path.exists(ssd_template):
            try:
                ssd_data = self.ssd_json(enhanced_data)
                if ssd_data:
                    doc2 = Document(ssd_template)
                    para_ssd = {
                        "{{PROJECTNAME}}": ssd_data['ProjectName'],
                        "{{PURPOSE1}}": ssd_data['Purpose'],
                        "{{SCOPE1}}": ssd_data['Scope'],
                        "{{VERSION}}": "1.0",
                        "{{DATE}}": today,
                        "{{AUTHOR}}": author
                    }
                    for k, v in para_ssd.items(): self._apply_paragraphs(k, v, doc2)
                    
                    bull_ssd = {
                        "{{SYSTEMARCHITECTURE}}": "SystemArchitecture",
                        "{{HARDWARECOMPONENTSPECIFICATION}}": "HardwareComponentSpecification",
                        "{{SOFTWARESTACK}}": "SoftwareStack",
                        "{{CONSTRAINTS}}": "Constraints",
                        "{{INTERFACES}}": "Interfaces",
                        "{{SAFETY&FAILOVER}}": "SafetyFailover",
                        "{{ELECTRICAL & MECHANICAL INTEGRATION NOTES}}": "ElectricalMechanicalIntegration"
                    }
                    for k, v in bull_ssd.items(): self._insert_bullets(k, ssd_data[v], doc2)
                    
                    doc2.save(ssd_path)
                    ssd_generated = True
                    print(f"SSD generated: {ssd_path}")
            except Exception as e:
                print(f"SSD generation error: {e}")
        else:
            print(f"SSD template missing at {ssd_template}")
            
        # Update DB only for what was generated
        if srd_generated or ssd_generated:
            if srd_generated and ssd_generated:
                conn.execute("UPDATE projects SET srd_path = ?, ssd_path = ? WHERE id = ?", (srd_filename, ssd_filename, project_id))
            elif srd_generated:
                conn.execute("UPDATE projects SET srd_path = ? WHERE id = ?", (srd_filename, project_id))
            elif ssd_generated:
                conn.execute("UPDATE projects SET ssd_path = ? WHERE id = ?", (ssd_filename, project_id))
            conn.commit()
        conn.close()

    def start_background_generation(self, project_id):
        thread = threading.Thread(target=self.generate_project_docs, args=(project_id,))
        thread.daemon = True
        thread.start()
