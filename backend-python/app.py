import os
import hashlib
import sqlite3
import json
import threading
import time
from datetime import datetime, timedelta
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import google.generativeai as genai
from docx import Document
import io

# Import existing AI utilities
import sys
sys.path.append(os.path.join(os.getcwd(), 'backend-python'))
from groq_client import call_groq
from meeting_utils import send_meeting_invite
from prompt import SYSTEM_PROMPT
from document_service import DocumentService

app = Flask(__name__)
CORS(app)

_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(_BASE_DIR, "tracker.db")
CONFIG_FILE = os.path.join(_BASE_DIR, "email_config.json")
UPLOADS_DIR = os.path.join(_BASE_DIR, "uploads")

doc_service = DocumentService(DB_PATH, CONFIG_FILE)

def get_db_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def load_email_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r") as f:
                return json.load(f)
        except:
            return {}
    return {}

def hash_password(pw: str) -> str:
    return hashlib.sha256(pw.encode()).hexdigest()

# ---------- AUTH & EMPLOYEES ----------

@app.route('/api/login', methods=['POST'])
def login():
    data = request.json
    email = data.get('email', '').strip().lower()
    password = data.get('password', '')
    
    print(f"DEBUG: Login attempt for email: {email}")
    
    conn = get_db_connection()
    user = conn.execute("SELECT * FROM employees WHERE LOWER(email)=?", (email,)).fetchone()
    conn.close()

    if not user:
        print(f"DEBUG: No user found with email: {email}")
    else:
        print(f"DEBUG: User found: {user['name']}, Role: {user['role']}")
        # Don't print password in production, but this is debugging
        # print(f"DEBUG: DB Pass: {user['password']}, Input Hash: {hash_password(password)}")
    
    if user:
        if password == user['password'] or hash_password(password) == user['password']:
            return jsonify({
                "role": user['role'],
                "name": user['name'],
                "email": user['email']
            })
    return jsonify({"error": "Invalid credentials"}), 401

@app.route('/api/users/employees', methods=['GET', 'POST'])
def handle_employees():
    conn = get_db_connection()
    if request.method == 'GET':
        employees = conn.execute("SELECT id, name, email, role, dob FROM employees").fetchall()
        conn.close()
        return jsonify([dict(emp) for emp in employees])
    
    if request.method == 'POST':
        data = request.json
        name = data.get('name')
        email = data.get('email')
        password = data.get('password', 'password123')
        role = data.get('role', 'Developer')
        dob = data.get('dob', '')
        
        try:
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO employees (name, email, password, role, dob) VALUES (?,?,?,?,?)",
                (name, email, hash_password(password), role, dob)
            )
            conn.commit()
            new_id = cursor.lastrowid
            conn.close()
            return jsonify({"user": {"id": new_id, "name": name, "email": email, "role": role, "dob": dob}}), 201
        except Exception as e:
            conn.close()
            return jsonify({"error": str(e)}), 400

@app.route('/api/users/employees/<int:id>', methods=['PUT', 'DELETE'])
def update_delete_employee(id):
    conn = get_db_connection()
    if request.method == 'DELETE':
        conn.execute("DELETE FROM employees WHERE id=?", (id,))
        conn.commit()
        conn.close()
        return jsonify({"success": True})
    
    if request.method == 'PUT':
        data = request.json
        name = data.get('name')
        email = data.get('email')
        role = data.get('role')
        dob = data.get('dob')
        username = data.get('username')
        password = data.get('password')
        
        try:
            if password and len(password.strip()) > 0:
                # Update with new password
                hashed_pw = hash_password(password)
                conn.execute(
                    "UPDATE employees SET name=?, email=?, role=?, dob=?, username=?, password=? WHERE id=?",
                    (name, email, role, dob, username, hashed_pw, id)
                )
            else:
                # Update without changing password
                conn.execute(
                    "UPDATE employees SET name=?, email=?, role=?, dob=?, username=? WHERE id=?",
                    (name, email, role, dob, username, id)
                )
            conn.commit()
            conn.close()
            return jsonify({"success": True})
        except Exception as e:
            conn.close()
            return jsonify({"error": str(e)}), 400

# ---------- PROJECTS ----------

@app.route('/api/projects', methods=['GET', 'POST'])
def handle_projects():
    conn = get_db_connection()
    if request.method == 'GET':
        projects_rows = conn.execute("SELECT * FROM projects").fetchall()
        projects = [dict(p) for p in projects_rows]
        
        # Fetch members and current milestone for each project
        for p in projects:
            members = conn.execute('''SELECT e.name, e.role 
                                     FROM employees e 
                                     JOIN project_members pm ON e.id = pm.employee_id 
                                     WHERE pm.project_id = ?''', (p['id'],)).fetchall()
            p['members'] = [dict(m) for m in members]
            
            # Highest milestone among tasks
            res = conn.execute("SELECT MAX(milestone) FROM tasks WHERE project_id = ?", (p['id'],)).fetchone()
            p['current_milestone'] = res[0] if res and res[0] else 0
            
        conn.close()
        return jsonify(projects)
    
    if request.method == 'POST':
        data = request.json
        name = data.get('name')
        owner = data.get('owner')
        email = data.get('email')
        description = data.get('description', '')
        station = data.get('station', '')
        category = data.get('category', '')
        itta = data.get('itta', '')
        member_ids = data.get('member_ids', [])
        
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO projects (name, owner, email, description, station, is_completed, category, itta) VALUES (?,?,?,?,?,?,?,?)",
            (name, owner, email, description, station, 0, category, itta)
        )
        project_id = cursor.lastrowid
        
        # Add members if provided
        for emp_id in member_ids:
            try:
                cursor.execute("INSERT INTO project_members (project_id, employee_id) VALUES (?,?)", (project_id, emp_id))
            except sqlite3.IntegrityError:
                pass
                
        conn.commit()
        log_activity_internal(project_id, f"Created project: {name}")
        
        # Trigger Document Generation
        doc_service.start_background_generation(project_id)
        
        conn.close()
        return jsonify({"success": True, "id": project_id}), 201

@app.route('/api/projects/<int:id>', methods=['PUT', 'DELETE'])
def update_delete_project(id):
    conn = get_db_connection()
    if request.method == 'DELETE':
        # Also cleanup project members
        conn.execute("DELETE FROM project_members WHERE project_id=?", (id,))
        conn.execute("DELETE FROM projects WHERE id=?", (id,))
        conn.commit()
        conn.close()
        return jsonify({"success": True})
    
    if request.method == 'PUT':
        data = request.json
        name = data.get('name')
        owner = data.get('owner')
        email = data.get('email')
        is_completed = data.get('is_completed', 0)
        description = data.get('description', '')
        station = data.get('station', '')
        category = data.get('category', '')
        itta = data.get('itta', '')
        
        conn.execute(
            "UPDATE projects SET name=?, owner=?, email=?, description=?, station=?, is_completed=?, category=?, itta=? WHERE id=?",
            (name, owner, email, description, station, is_completed, category, itta, id)
        )
        conn.commit()
        conn.close()
        return jsonify({"success": True})

# ---------- TASKS ----------

@app.route('/api/tasks', methods=['GET', 'POST'])
def handle_tasks():
    conn = get_db_connection()
    if request.method == 'GET':
        project_id = request.args.get('project_id')
        if project_id:
            tasks = conn.execute("SELECT * FROM tasks WHERE project_id=?", (project_id,)).fetchall()
        else:
            tasks = conn.execute("SELECT * FROM tasks").fetchall()
        conn.close()
        return jsonify([dict(t) for t in tasks])
    
    if request.method == 'POST':
        data = request.json
        project_id = data.get('project_id')
        title = data.get('title')
        due_date = data.get('dueDate')
        priority = data.get('priority')
        assignee = data.get('assignee', '')
        created_by = data.get('created_by', '')
        
        # Fetch assignee email for storage
        assignee_email = None
        if assignee:
            conn_internal = sqlite3.connect(DB_PATH)
            emp = conn_internal.execute("SELECT email FROM employees WHERE name = ?", (assignee,)).fetchone()
            conn_internal.close()
            if emp:
                assignee_email = emp[0]

        milestone = data.get('milestone') # Expected 1-8

        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO tasks (task, due_date, priority, assignee, assignee_email, progress, done, project_id, created_by, milestone) 
            VALUES (?,?,?,?,?,?,?,?,?,?)
        """, (title, due_date, priority, assignee, assignee_email, 0, 0, project_id, created_by, milestone))
        task_id = cursor.lastrowid
        conn.commit()
        conn.close()

        # Send Assignment Email
        if assignee_email:
            subject = f"New Task Assigned: {title}"
            body = f"Hello,\n\nA new task has been assigned to you.\n\nTask: {title}\nDue Date: {due_date}\nPriority: {priority}\n\nGood luck!\nAI PM Team"
            send_email(assignee_email, subject, body)

        log_activity_internal(project_id, f"New task created: {title} (Assigned to: {assignee})")
        return jsonify({"message": "Task created and notification sent", "id": task_id}), 201

@app.route('/api/tasks/<int:id>', methods=['PUT', 'DELETE'])
def update_delete_task(id):
    conn = get_db_connection()
    if request.method == 'DELETE':
        conn.execute("DELETE FROM tasks WHERE id=?", (id,))
        conn.commit()
        conn.close()
        return jsonify({"success": True})
    
    if request.method == 'PUT':
        data = request.json
        progress = data.get('progress', 0)
        done = 1 if progress == 100 else 0
        
        # Fetch current task state to check for completion transition
        old_task = conn.execute("SELECT * FROM tasks WHERE id = ?", (id,)).fetchone()
        
        conn.execute("UPDATE tasks SET progress=?, done=? WHERE id=?", (progress, done, id))
        conn.commit()

        # Send Completion Email if transition to 100% occurs
        old_progress = old_task['progress'] if old_task['progress'] is not None else 0
        if progress == 100 and old_progress < 100:
            project = conn.execute("SELECT * FROM projects WHERE id = ?", (old_task['project_id'],)).fetchone()
            
            # Determine recipient: Project Email > Cached Assignee Email > DB Assignee Email
            recipient = None
            if project and project['email']:
                recipient = project['email']
            elif old_task['assignee_email']:
                recipient = old_task['assignee_email']
            else:
                # Last resort: lookup in employees table
                emp = conn.execute("SELECT email FROM employees WHERE name = ?", (old_task['assignee'],)).fetchone()
                if emp:
                    recipient = emp[0]
            
            if recipient:
                project_name = project['name'] if project else "General"
                subject = f"✅ Task Completed: {old_task['task']}"
                body = f"Hello,\n\nThe task '{old_task['task']}' in project '{project_name}' has been marked as COMPLETED by {old_task['assignee']}.\n\nAI PM Team"
                send_email(recipient, subject, body)
                log_activity_internal(old_task['project_id'], f"Task completed: {old_task['task']} (Assignee: {old_task['assignee']})")
                print(f"Completion email sent to {recipient}")
            else:
                print(f"⚠️ Skipping completion email for task {id}: No recipient found.")

        conn.close()
        return jsonify({"success": True})

def due_date_monitor():
    """Background thread to monitor due dates and send notifications."""
    print("Due date monitor thread started")
    while True:
        try:
            conn = sqlite3.connect(DB_PATH)
            conn.row_factory = sqlite3.Row
            today = datetime.now().strftime("%Y-%m-%d")
            
            # --- 1. Overdue Alerts ---
            # Use LEFT JOIN to catch personal tasks without projects
            overdue_tasks = conn.execute("""
                SELECT t.*, p.name as project_name, COALESCE(t.assignee_email, e.email) as final_email
                FROM tasks t
                LEFT JOIN projects p ON t.project_id = p.id
                LEFT JOIN employees e ON LOWER(TRIM(t.assignee)) = LOWER(TRIM(e.name))
                WHERE t.due_date < ? AND t.done = 0 AND t.progress < 100 AND (t.alert_sent = 0 OR t.alert_sent IS NULL)
            """, (today,)).fetchall()
            
            if overdue_tasks:
                print(f"Found {len(overdue_tasks)} overdue tasks to notify.")
            
            for task in overdue_tasks:
                if not task['final_email']:
                    print(f"⚠️ No email found for overdue task: {task['task']} (Assignee: {task['assignee']})")
                    continue
                
                project_label = task['project_name'] if task['project_name'] else "General"
                subject = f"⚠️ OVERDUE TASK: {task['task']}"
                body = f"Hello {task['assignee']},\n\nYour task '{task['task']}' in project '{project_label}' is OVERDUE (Due: {task['due_date']}).\n\nPlease update the status as soon as possible.\n\nAI PM Team"
                if send_email(task['final_email'], subject, body):
                    conn.execute("UPDATE tasks SET alert_sent = 1 WHERE id = ?", (task['id'],))
            
            # --- 2. Today Reminders ---
            today_tasks = conn.execute("""
                SELECT t.*, p.name as project_name, COALESCE(t.assignee_email, e.email) as final_email
                FROM tasks t
                LEFT JOIN projects p ON t.project_id = p.id
                LEFT JOIN employees e ON LOWER(TRIM(t.assignee)) = LOWER(TRIM(e.name))
                WHERE t.due_date = ? AND t.done = 0 AND t.progress < 100 AND (t.reminder_sent = 0 OR t.reminder_sent IS NULL)
            """, (today,)).fetchall()
            
            if today_tasks:
                print(f"Found {len(today_tasks)} tasks due today to notify.")
            
            for task in today_tasks:
                if not task['final_email']:
                    print(f"⚠️ No email found for today reminder: {task['task']} (Assignee: {task['assignee']})")
                    continue
                
                project_label = task['project_name'] if task['project_name'] else "General"
                subject = f"📅 Reminder: Task Due Today - {task['task']}"
                body = f"Hello {task['assignee']},\n\nThis is a reminder that your task '{task['task']}' in project '{project_label}' is due TODAY ({task['due_date']}).\n\nGood luck!\nAI PM Team"
                if send_email(task['final_email'], subject, body):
                    conn.execute("UPDATE tasks SET reminder_sent = 1 WHERE id = ?", (task['id'],))
            
            conn.commit()
            conn.close()
        except Exception as e:
            print(f"⚠️ Monitor Error: {e}")
        
        # Check every 5 minutes
        time.sleep(300)

# monitor_thread = threading.Thread(target=due_date_monitor, daemon=True)
# monitor_thread.start()

# ---------- DASHBOARD & ACTIVITY ----------

@app.route('/api/stats/admin', methods=['GET'])
def admin_stats():
    station = request.args.get('station')
    category = request.args.get('category')
    itta = request.args.get('itta')
    
    conn = get_db_connection()
    
    # 1. Projects Filter Logic
    proj_query = "SELECT * FROM projects WHERE 1=1"
    proj_params = []
    if station:
        proj_query += " AND station = ?"
        proj_params.append(station)
    if category:
        proj_query += " AND category = ?"
        proj_params.append(category)
    if itta:
        proj_query += " AND itta = ?"
        proj_params.append(itta)
        
    projects = conn.execute(proj_query, proj_params).fetchall()
    project_ids = [p['id'] for p in projects]
    
    # Project Breakdown
    active_projects = len([p for p in projects if p['is_completed'] == 0])
    completed_projects = len([p for p in projects if p['is_completed'] == 1])
    
    # 2. Tasks Filter Logic
    if project_ids:
        placeholders = ','.join(['?'] * len(project_ids))
        tasks = conn.execute(f"SELECT * FROM tasks WHERE project_id IN ({placeholders})", project_ids).fetchall()
    else:
        # If no projects match filters, stats should be zero
        tasks = []
        
    total_tasks = len(tasks)
    done_tasks = len([t for t in tasks if t['done'] == 1])
    pending_tasks = total_tasks - done_tasks
    
    # Category Breakdown
    cat_counts = {"Forge Labs Project": 0, "LightHouse Project": 0, "Startup Project": 0}
    for p in projects:
        c = p['category']
        if c in cat_counts:
            cat_counts[c] += 1
            
    # ITTA Breakdown
    itta_counts = {"HSE": 0, "AMDF": 0, "AMR": 0, "ACCS": 0, "IIDP": 0}
    for p in projects:
        i = p['itta']
        if i in itta_counts:
            itta_counts[i] += 1
            
    # Employee count
    employee_count = conn.execute("SELECT COUNT(*) FROM employees").fetchone()[0]

    # On-Hold Logic: Active projects with no activity for > 7 days
    on_hold_projects = 0
    seven_days_ago = (datetime.now() - timedelta(days=7)).strftime("%Y-%m-%d %H:%M:%S")
    
    for p in projects:
        if p['is_completed'] == 0:
            last_act = conn.execute("SELECT MAX(time) FROM activity_log WHERE project_id = ?", (p['id'],)).fetchone()[0]
            if not last_act or last_act < seven_days_ago:
                on_hold_projects += 1
                
    conn.close()
    return jsonify({
        "employees": employee_count,
        "projects": len(projects),
        "active_projects": active_projects,
        "completed_projects": completed_projects,
        "on_hold_projects": on_hold_projects,
        "completed_tasks": done_tasks,
        "pending_tasks": pending_tasks,
        "total_tasks": total_tasks,
        "category_stats": [
            {"name": "Forge Labs", "value": cat_counts["Forge Labs Project"]},
            {"name": "LightHouse", "value": cat_counts["LightHouse Project"]},
            {"name": "Startup", "value": cat_counts["Startup Project"]}
        ],
        "itta_stats": [
            {"name": "HSE", "value": itta_counts["HSE"]},
            {"name": "AMDF", "value": itta_counts["AMDF"]},
            {"name": "AMR", "value": itta_counts["AMR"]},
            {"name": "ACCS", "value": itta_counts["ACCS"]},
            {"name": "IIDP", "value": itta_counts["IIDP"]}
        ]
    })

@app.route('/api/stats/employee', methods=['GET'])
def employee_stats():
    email = request.args.get('email')
    name = request.args.get('name')
    if not email:
        return jsonify({"error": "Email required"}), 400
        
    conn = get_db_connection()
    # Fetch tasks WHERE user is assignee OR assignee_email
    query = """
        SELECT * FROM tasks 
        WHERE LOWER(assignee) = ? 
           OR LOWER(assignee) = ? 
           OR LOWER(assignee_email) = ?
    """
    params = (
        email.lower(), 
        name.lower() if name else email.lower(), 
        email.lower()
    )
    tasks = conn.execute(query, params).fetchall()
    
    assigned = len(tasks)
    completed = len([t for t in tasks if t['done'] == 1 or t['progress'] == 100])
    pending = len([t for t in tasks if (t['done'] == 0 or t['done'] is None) and t['progress'] < 100])
    
    conn.close()
    return jsonify({
        "assigned": assigned,
        "completed": completed,
        "pending": pending,
        "tasks": [dict(t) for t in tasks]
    })

# ---------- ACTIVITY ----------

@app.route('/api/activity', methods=['GET'])
def get_activity():
    conn = get_db_connection()
    logs = conn.execute("SELECT * FROM activity_log ORDER BY time DESC LIMIT 50").fetchall()
    conn.close()
    return jsonify([dict(log) for log in logs])

# ---------- DOCUMENTS ----------

@app.route('/api/projects/documents/<filename>')
def serve_document(filename):
    file_path = os.path.join(UPLOADS_DIR, filename)
    if not os.path.exists(file_path):
        return jsonify({"error": "File not found"}), 404
    return send_file(file_path, as_attachment=True)

@app.route('/api/projects/<int:id>/generate-docs', methods=['POST'])
def manual_generate_docs(id):
    doc_service.start_background_generation(id)
    return jsonify({"success": True, "message": "Document generation started"})

# ---------- PROJECT MEMBERS ----------

@app.route('/api/projects/<int:id>/members', methods=['GET', 'POST'])
def project_members(id):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if request.method == 'GET':
        cursor.execute('''SELECT e.id, e.name, e.email, e.role 
                         FROM employees e 
                         JOIN project_members pm ON e.id = pm.employee_id 
                         WHERE pm.project_id = ?''', (id,))
        members = [dict(zip(['id', 'name', 'email', 'role'], row)) for row in cursor.fetchall()]
        conn.close()
        return jsonify(members)
    
    if request.method == 'POST':
        data = request.json
        emp_id = data.get('employee_id')
        try:
            cursor.execute("INSERT INTO project_members (project_id, employee_id) VALUES (?,?)", (id, emp_id))
            conn.commit()
            conn.close()
            return jsonify({"message": "Member added to project"}), 201
        except sqlite3.IntegrityError:
            conn.close()
            return jsonify({"error": "Member already in project"}), 400

@app.route('/api/projects/<int:proj_id>/members/<int:emp_id>', methods=['DELETE'])
def remove_project_member(proj_id, emp_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM project_members WHERE project_id = ? AND employee_id = ?", (proj_id, emp_id))
    conn.commit()
    conn.close()
    return jsonify({"message": "Member removed from project"})

# ---------- AI ENDPOINTS ----------

@app.route('/api/ai/generate-doc', methods=['POST'])
def generate_doc():
    data = request.json
    description = data.get('description', '')
    config = load_email_config()
    api_key = config.get("groq_key", "")
    
    prompt = f"Write a professional document based on this description: {description}. Use clear headings and sections."
    content = call_groq(prompt, api_key)
    return jsonify({"content": content})

@app.route('/api/ai/export-word', methods=['POST'])
def export_word():
    data = request.json
    title = data.get('title', 'Document')
    content = data.get('content', '')
    
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    
    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    
    return send_file(
        target,
        as_attachment=True,
        download_name=f"{title.replace(' ', '_')}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/api/ai/task-intent', methods=['POST'])
def task_intent():
    data = request.json
    user_input = data.get("text", "")
    config = load_email_config()
    api_key = config.get("groq_key", "")
    
    prompt = f"{SYSTEM_PROMPT}\n\nUser input: {user_input}"
    ai_response = call_groq(prompt, api_key)
    
    try:
        parsed = json.loads(ai_response)
        return jsonify(parsed)
    except:
        return jsonify({"title": user_input, "priority": "Medium", "assignee": "Unassigned"})

@app.route('/api/ai/chat', methods=['POST'])
def ai_chat():
    data = request.json
    message = data.get("message", "")
    config = load_email_config()
    api_key = config.get("groq_key", "")
    
    prompt = f"You are a helpful project management assistant. User says: {message}"
    response = call_groq(prompt, api_key)
    return jsonify({"reply": response})

@app.route('/api/ai/smart-chat', methods=['POST'])
def smart_chat():
    """
    Full-context AI endpoint. Injects live data (employees, projects, tasks)
    so the AI can resolve names accurately and return structured intents.
    """
    data = request.json
    message = data.get("message", "")
    user_email = data.get("user_email", "")
    user_name = data.get("user_name", "")
    user_role = data.get("user_role", "")

    config = load_email_config()
    api_key = config.get("groq_key", "")

    conn = get_db_connection()

    # Fetch live context
    employees = conn.execute("SELECT id, name, email, role FROM employees").fetchall()
    projects = conn.execute("SELECT id, name, description FROM projects").fetchall()

    # Tasks relevant to the user (all tasks for admin, own tasks for employee)
    if user_role and user_role.lower() == 'admin':
        tasks = conn.execute("SELECT id, task, assignee, progress, done FROM tasks LIMIT 50").fetchall()
    else:
        tasks = conn.execute(
            "SELECT id, task, assignee, progress, done FROM tasks WHERE LOWER(assignee) = ? OR LOWER(assignee_email) = ? LIMIT 30",
            (user_name.lower(), user_email.lower())
        ).fetchall()
    conn.close()

    emp_list = [f"{e['name']} ({e['role']}, {e['email']})" for e in employees]
    proj_list = [f"[ID:{p['id']}] {p['name']}" for p in projects]
    def task_status(t):
        if t['done']: return 'done'
        if t['progress'] and t['progress'] > 0: return 'inprogress'
        return 'todo'
    task_list = [f"[ID:{t['id']}] '{t['task']}' assigned to {t['assignee']} | status:{task_status(t)} | progress:{t['progress']}%" for t in tasks]

    context_block = f"""
CONTEXT (use this for resolving names):
Current User: {user_name} ({user_role}, {user_email})
Today: {datetime.now().strftime('%Y-%m-%d')}

Employees:
{chr(10).join(emp_list) or 'None'}

Projects:
{chr(10).join(proj_list) or 'None'}

Tasks:
{chr(10).join(task_list) or 'None'}
"""

    full_prompt = f"{SYSTEM_PROMPT}\n\n{context_block}\n\nUser message: {message}"
    ai_response = call_groq(full_prompt, api_key)

    try:
        parsed = json.loads(ai_response)
        # Build a human reply alongside the structured intent
        action = parsed.get("action", "chat")
        if action == "add_task":
            assignee = parsed.get("assignee") or user_name
            reply = f"Got it! Creating task **\"{parsed.get('title')}\"** and assigning it to **{assignee if assignee != 'SELF' else user_name}**."
        elif action == "create_project":
            reply = f"Creating project **\"{parsed.get('name')}\"** now!"
        elif action == "update_task_status":
            reply = f"Updating **\"{parsed.get('task_title')}\"** to **{parsed.get('status') or 'done'}**."
        elif action == "delete_task":
            reply = f"Deleting task **\"{parsed.get('task_title')}\"**."
        elif action == "list_tasks":
            if task_list:
                reply = "Here are the tasks:\n" + "\n".join([f"• {t['task']} → {t['assignee']} ({task_status(t)})" for t in tasks])
            else:
                reply = "No tasks found."
        elif action == "list_projects":
            if proj_list:
                reply = "Current projects:\n" + "\n".join([f"• {p['name']}" for p in projects])
            else:
                reply = "No projects found."
        elif action == "add_member":
            reply = f"Adding **{parsed.get('member_name')}** to project **{parsed.get('project_name')}**."
        else:
            # Fall back to conversational for 'chat' action
            conv_prompt = f"""You are a friendly and helpful project management assistant named Alex. 
Answer the following question helpfully and concisely. You can tell the user what you are capable of doing:
- Create tasks, projects, assign team members
- Update task status (todo/in progress/done)
- List tasks and projects
- Schedule meetings

User: {message}"""
            reply = call_groq(conv_prompt, api_key)

        parsed["reply"] = reply
        parsed["tasks_context"] = [dict(t) for t in tasks] if action in ("list_tasks",) else []
        parsed["projects_context"] = [dict(p) for p in projects] if action in ("list_projects",) else []
        return jsonify(parsed)

    except Exception as e:
        # Full conversational fallback
        conv_prompt = f"You are a friendly project management assistant named Alex. Answer this helpfully: {message}"
        reply = call_groq(conv_prompt, api_key)
        return jsonify({"action": "chat", "reply": reply})


def log_activity_internal(project_id, message):
    conn = get_db_connection()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn.execute("INSERT INTO activity_log (project_id, message, time) VALUES (?,?,?)", (project_id, message, now))
    conn.commit()
    conn.close()

import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

def send_email(to_email, subject, body):
    if not to_email:
        print("⚠️ Recipient email is empty, skipping email.")
        return False
        
    config = load_email_config()
    sender = config.get("sender_email")
    password = config.get("app_password")
    
    if not sender or not password:
        print("⚠️ Email credentials missing in email_config.json")
        return False

    try:
        msg = MIMEMultipart()
        msg['From'] = str(sender)
        msg['To'] = str(to_email)
        msg['Subject'] = str(subject)
        msg.attach(MIMEText(body, 'plain'))

        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(str(sender), str(password))
        server.send_message(msg)
        server.quit()
        print(f"Email sent to {to_email}")
        return True
    except Exception as e:
        print(f"Failed to send email: {e}")
        return False

def init_db():
    conn = sqlite3.connect('tracker.db')
    c = conn.cursor()
    # Create tables
    c.execute('''CREATE TABLE IF NOT EXISTS employees 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, email TEXT, password TEXT, role TEXT, dob TEXT, username TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS projects 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, owner TEXT, email TEXT, description TEXT, station TEXT, category TEXT, itta TEXT, srd_path TEXT, ssd_path TEXT, is_completed INTEGER DEFAULT 0)''')
    
    # Migration: Check for Category, Purpose, ITTA columns
    cursor = c.execute("PRAGMA table_info(projects)")
    columns = [row[1] for row in cursor.fetchall()]
    if 'category' not in columns:
        c.execute("ALTER TABLE projects ADD COLUMN category TEXT DEFAULT ''")
    if 'itta' not in columns:
        c.execute("ALTER TABLE projects ADD COLUMN itta TEXT DEFAULT ''")
    if 'srd_path' not in columns:
        c.execute("ALTER TABLE projects ADD COLUMN srd_path TEXT DEFAULT ''")
    if 'ssd_path' not in columns:
        c.execute("ALTER TABLE projects ADD COLUMN ssd_path TEXT DEFAULT ''")
    
    c.execute('''CREATE TABLE IF NOT EXISTS tasks 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, task TEXT, due_date TEXT, priority TEXT, assignee TEXT, assignee_email TEXT, progress INTEGER, done INTEGER, project_id INTEGER, alert_sent INTEGER, reminder_sent INTEGER, reminded INTEGER, created_by TEXT, milestone INTEGER)''')
    c.execute('''CREATE TABLE IF NOT EXISTS activity_log 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, project_id INTEGER, message TEXT, time TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS project_members 
                 (project_id INTEGER, employee_id INTEGER, PRIMARY KEY(project_id, employee_id))''')
    
    # Migration for milestone
    cursor = c.execute("PRAGMA table_info(tasks)")
    columns = [row[1] for row in cursor.fetchall()]
    if 'milestone' not in columns:
        c.execute("ALTER TABLE tasks ADD COLUMN milestone INTEGER")
    
    # Seed a default admin if no users exist
    existing = c.execute("SELECT COUNT(*) FROM employees").fetchone()[0]
    if existing == 0:
        import hashlib
        pw_hash = hashlib.sha256("admin123".encode()).hexdigest()
        c.execute(
            "INSERT INTO employees (name, email, password, role, username) VALUES (?,?,?,?,?)",
            ("Admin User", "admin123", pw_hash, "Admin", "admin123")
        )
        conn.commit()
        print("Default admin created: admin123 / admin123")
    conn.close()

if __name__ == '__main__':
    init_db()

    monitor_thread = threading.Thread(target=due_date_monitor, daemon=True)
    monitor_thread.start()

    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
